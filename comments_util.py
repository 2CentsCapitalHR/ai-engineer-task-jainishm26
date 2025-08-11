from docx import Document
from docx.text.run import Run
from lxml import etree
from docx.oxml.shared import qn
from datetime import datetime

def _add_comment(part, paragraph, text, author="ADGM Agent"):
    """
    Insert a Word comment (OOXML) at the start of `paragraph`.
    Falls back to text highlight if XML insertion fails.
    """
    try:
        doc_part = part
        comments_part = None
        for rel in doc_part.part.rels.values():
            if "comments" in rel.reltype:
                comments_part = rel._target
                break
        if comments_part is None:
            comments_part = doc_part.part.add_comments_part()
            comments_part._element.append(etree.Element(qn('w:comments')))

        comments_el = comments_part._element
        # build comment
        cmt_id = str(int(datetime.now().timestamp()))
        cmt = etree.Element(qn('w:comment'), {qn('w:id'): cmt_id})
        p = etree.SubElement(cmt, qn('w:p'))
        r = etree.SubElement(p, qn('w:r'))
        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        comments_el.append(cmt)

        # mark range in body
        p_el = paragraph._p
        start = etree.Element(qn('w:commentRangeStart'), {qn('w:id'): cmt_id})
        end = etree.Element(qn('w:commentRangeEnd'), {qn('w:id'): cmt_id})
        p_el.addprevious(start)
        p_el.addnext(end)
        # add reference
        r_ref = etree.Element(qn('w:r'))
        rpr = etree.SubElement(r_ref, qn('w:rPr'))
        ar = etree.SubElement(rpr, qn('w:rStyle'))
        ar.set(qn('w:val'), "CommentReference")
        cref = etree.SubElement(r_ref, qn('w:commentReference'), {qn('w:id'): cmt_id})
        p_el.append(r_ref)
        return True
    except Exception:
        # fallback: highlight paragraph and append inline note
        run: Run = paragraph.add_run(f" [COMMENT] {text}")
        run.font.highlight_color = 3  # yellow
        return False

def add_comment_at_paragraph(doc: Document, para_idx: int, text: str):
    if para_idx < 0 or para_idx >= len(doc.paragraphs):
        return False
    paragraph = doc.paragraphs[para_idx]
    return _add_comment(doc, paragraph, text)
