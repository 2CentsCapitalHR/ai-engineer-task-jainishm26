import json
from pathlib import Path
from typing import List, Tuple, Dict, Any
from docx import Document
from checklist import classify_doc_type, REQUIRED_DOCS
from comment_utils import add_comment_at_paragraph
from utils import ensure_dirs

# Simple rule patterns 
BAD_JURISDICTION = ["uae federal court", "dubai courts", "onshore uae"]
WEAK_LANGUAGE = ["may at its discretion", "best efforts", "commercially reasonable efforts"]
MISSING_SIGN_KEYS = ["signature", "signed by", "authorised signatory", "authorized signatory", "date", "name"]

def _extract_text(doc_path: str, max_chars=20000) -> str:
    d = Document(doc_path)
    text = "\n".join(p.text for p in d.paragraphs)
    return text[:max_chars]

def _find_red_flags(text: str, retriever, doc_type: str) -> List[Dict[str, Any]]:
    lowers = text.lower()
    issues = []

    # 1) Wrong jurisdiction
    for term in BAD_JURISDICTION:
        if term in lowers:
            refs = retriever.get_relevant_documents("ADGM jurisdiction clause Companies Regulations courts venue")
            citation = refs[0].metadata.get("source", "ADGM Regulation") if refs else "ADGM Regulation"
            issues.append({
                "issue": "Jurisdiction references onshore UAE/Federal Courts",
                "severity": "High",
                "suggestion": "Update governing law and forum to ADGM Courts.",
                "citation": citation
            })
            break

    # 2) Weak language
    for term in WEAK_LANGUAGE:
        if term in lowers:
            refs = retriever.get_relevant_documents(f"binding language {doc_type} ADGM template clause shall must")
            citation = refs[0].metadata.get("source", "ADGM Guidance/Template") if refs else "ADGM Guidance/Template"
            issues.append({
                "issue": f"Ambiguous/weak obligation: '{term}'",
                "severity": "Medium",
                "suggestion": "Prefer firm language ('shall', specific obligations).",
                "citation": citation
            })

    # 3) Signing block present?
    if not any(k in lowers for k in MISSING_SIGN_KEYS):
        refs = retriever.get_relevant_documents(f"{doc_type} signature block ADGM template")
        citation = refs[0].metadata.get("source", "ADGM Template") if refs else "ADGM Template"
        issues.append({
            "issue": "Missing or incomplete signatory section.",
            "severity": "High",
            "suggestion": "Add an authorised signatory block with name, title, date, signature.",
            "citation": citation
        })

    return issues

def _insert_comments(doc_path: str, issues: List[Dict[str, Any]]) -> str:
    doc = Document(doc_path)
    note_lines = []
    for i, it in enumerate(issues):
        note = f"{it['issue']} | Suggestion: {it['suggestion']} | Source: {it['citation']}"
        note_lines.append(note)
        # Add a comment near the first non-empty paragraph
        target_idx = 0
        for idx, p in enumerate(doc.paragraphs):
            if p.text and len(p.text.strip()) > 3:
                target_idx = idx
                break
        add_comment_at_paragraph(doc, target_idx, note)
    out_name = Path(doc_path).stem + "_REVIEWED.docx"
    out_path = Path("outputs/reviewed") / out_name
    doc.save(out_path)
    return str(out_path)

def analyze_documents(
    uploaded_files: List[Tuple[str, str]],
    retriever,
    process_hint: str
) -> Dict[str, Any]:
    ensure_dirs()

    # Classify doc type + extract text
    doc_infos = []
    for orig_name, path in uploaded_files:
        text = _extract_text(path)
        dtype = classify_doc_type(orig_name, text)
        doc_infos.append({"name": orig_name, "path": path, "type": dtype, "text": text})

    process = process_hint
    required = REQUIRED_DOCS.get(process, [])
    present_types = {d["type"] for d in doc_infos if d["type"] != "Unknown"}
    missing = [d for d in required if d not in present_types]

    issues_found = []
    reviewed_paths = []
    for info in doc_infos:
        issues = _find_red_flags(info["text"], retriever, info["type"])
        if issues:
            reviewed_path = _insert_comments(info["path"], issues)
            reviewed_paths.append({"original_name": info["name"], "path": reviewed_path})
            for it in issues:
                issues_found.append({
                    "document": info["type"] if info["type"] != "Unknown" else info["name"],
                    "section": "N/A",
                    "issue": it["issue"],
                    "severity": it["severity"],
                    "suggestion": it["suggestion"],
                    "citation": it["citation"]
                })

    # Build structured report
    report = {
        "process": process,
        "documents_uploaded": len(uploaded_files),
        "required_documents": len(required),
        "missing_documents": missing,
        "issues_found": issues_found
    }
    report_path = Path("outputs/reports") / "report.json"
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2)

    report["reviewed_paths"] = reviewed_paths
    report["report_path"] = str(report_path)
    return report
